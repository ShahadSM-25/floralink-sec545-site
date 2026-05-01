from pathlib import Path
from docx import Document

DOC_PATH = Path('/home/ubuntu/upload/FloraLink_Final_Report_اساسي.docx')
doc = Document(str(DOC_PATH))

paragraph_targets = list(range(430, 461))
print('PARAGRAPHS')
for i in paragraph_targets:
    if i < len(doc.paragraphs):
        text = ' '.join(doc.paragraphs[i].text.split())
        if text:
            print(f'P#{i}: {text}')

print('\nTABLES')
for ti in [63, 64, 65, 66, 67, 68, 69]:
    if ti < len(doc.tables):
        print(f'-- TABLE {ti} --')
        table = doc.tables[ti]
        for ri, row in enumerate(table.rows):
            vals = [' '.join(cell.text.split()) for cell in row.cells]
            print(f'R{ri}: {vals}')
