from pathlib import Path
from docx import Document

SRC = Path('/home/ubuntu/upload/FloraLink_Final_Report_اساسي.docx')
OUT = Path('/home/ubuntu/floralink-sec545-site/FloraLink_Final_Report_corrected.docx')

doc = Document(str(SRC))


def replace_paragraph_containing(snippet: str, new_text: str):
    for paragraph in doc.paragraphs:
        if snippet in paragraph.text:
            paragraph.text = new_text
            return True
    raise ValueError(f'Paragraph containing snippet not found: {snippet}')


# Abstract / summary updates
replace_paragraph_containing(
    'Key results include a fully designed system architecture',
    'Key results include a fully designed system architecture, high-fidelity UI prototypes, a complete database design, managed logo integration, and a working full-stack authentication prototype. The implemented prototype now demonstrates registration, database-backed login, password reset, strong password validation, and temporary lockout behavior. Verification combined browser-based flow checks with automated backend route tests to confirm that the implemented authentication features behave correctly under valid and invalid inputs.'
)

# Implementation section
replace_paragraph_containing(
    'The following four items were implemented as a working prototype',
    'The working prototype currently centers on the FloraLink authentication flow. The original four deliverable items remain visible in the interface as the main functional scope, while the final implementation was extended with persistent database-backed storage and a forgot-password/reset-password flow to make the prototype more realistic and usable.'
)
replace_paragraph_containing(
    'UC-01 – Register Account:',
    'UC-01 – Register Account: The registration form collects full name, email, phone number, password, and confirmation. Before acceptance, the system checks that no required field is empty, validates the email format, compares the password against the confirmation field, enforces the password policy, and verifies through the backend that the email is not already registered. When registration succeeds, the account is inserted into the customer_accounts table and the interface displays a success screen with the saved account details.'
)
replace_paragraph_containing(
    'UC-02 – Login:',
    'UC-02 – Login: The login form no longer compares credentials against a hardcoded demo account. Instead, it sends the submitted email and password to the backend authentication procedure, which retrieves the corresponding customer record from the database and verifies the stored password hash. Successful login displays a welcome screen, while invalid credentials produce a generic error message and increment the failed-attempt counter in the interface.'
)
replace_paragraph_containing(
    'MIT-01 – Strong Authentication:',
    'MIT-01 – Strong Authentication: A real-time password strength indicator and interactive checklist enforce five rules: minimum 8 characters, at least one uppercase letter, at least one number, at least one special character, and rejection of common passwords. The same password rules are also validated in the backend procedures through schema validation so weak passwords are not accepted even if the frontend is bypassed.'
)
replace_paragraph_containing(
    'MIT-02 – Rate Limiting & Account Lockout:',
    'MIT-02 – Rate Limiting & Account Lockout: A warning is displayed on the fourth failed attempt, and on the fifth the interface places the session in a temporary 15-minute lock state with a visible countdown timer. The current version keeps this lockout behavior in the client-side interface for demonstration purposes, while the actual account credentials themselves are stored and verified through the database-backed backend.'
)
replace_paragraph_containing(
    '[Insert screenshots here:',
    '[Insert screenshots here: Registration Form, Registration Success, Login Form, Forgot Password / Reset Password Form, Password Updated Screen, Invalid Credentials Error, Pre-Lockout Warning (4th attempt), Account Locked Screen with Countdown, Password Strength Indicator]'
)

# Testing section
replace_paragraph_containing(
    'Tests were conducted on the locally running FloraLink web application served via Vite in development mode',
    'Tests were conducted on the locally running FloraLink web application through the current Vite + Express development setup, then rechecked in the browser preview after the backend/database upgrade. The environment uses the configured database connection and the customer_accounts table created through Drizzle migrations. Verification combined browser-based end-to-end checks for registration, login, and password reset with automated Vitest tests for the backend register, login, and reset-password procedures. Payment and notification services were outside the scope of this verification cycle.'
)
replace_paragraph_containing(
    'Web browser (Chrome/Firefox) – primary test execution interface.',
    'Web browser preview (Chromium) – used for end-to-end validation of registration, login, lock-state behavior, and password reset flows.'
)
replace_paragraph_containing(
    'Browser Developer Tools – for inspecting HTTP requests and interface state.',
    'Vitest + pnpm test – used to run repeatable automated tests for the backend registration, login, and reset-password routes.'
)
replace_paragraph_containing(
    'Manual test execution – each test case run manually and results recorded.',
    'Project verification commands (pnpm check, pnpm build, pnpm db:push) – used to confirm TypeScript integrity, production build readiness, and successful database schema synchronization.'
)
replace_paragraph_containing(
    'All 24 defined test cases were executed against the current FloraLink prototype.',
    'The report defines a 24-case black-box matrix for the four Phase 3 items. In the current verified build, the core user flows were rechecked in the browser after the database upgrade, and the backend credential routes were also covered by automated Vitest tests. These checks confirmed that the implemented authentication features work correctly in the current source state, although the full 24-case matrix should still be maintained as the formal manual test set for coursework reporting.'
)
replace_paragraph_containing(
    '[Fill in Pass/Fail counts after executing test cases]',
    'The detailed browser verification confirmed successful registration, successful login, successful password reset, and successful login with the updated password. Automated route tests also passed for successful registration, duplicate-email rejection, successful login, invalid login, successful password reset, and missing-account reset attempts.'
)
replace_paragraph_containing(
    'In the current prototype, the following limitations exist',
    'In the current prototype, the following implementation limitations and follow-up improvements remain relevant:'
)
replace_paragraph_containing(
    'No persistent database: account data is stored in memory only; a server restart resets all state.',
    'Customer account data is now stored persistently in the customer_accounts database table rather than in memory only.'
)
replace_paragraph_containing(
    'Demo credentials are hardcoded in the frontend code, which would be unacceptable in production.',
    'Passwords are no longer handled as hardcoded demo credentials; they are stored in the database as salted password hashes and verified on the backend.'
)
replace_paragraph_containing(
    'CAPTCHA is simulated rather than integrated with a real service (e.g., Google reCAPTCHA).',
    'The human-verification step is still a simulated UI element rather than a live CAPTCHA integration such as Google reCAPTCHA.'
)
replace_paragraph_containing(
    'Rate limiting is implemented in client-side state only; a real implementation requires server-side enforcement.',
    'The temporary lockout/rate-limiting behavior is still enforced in client-side state for demonstration purposes; a production deployment should mirror this control on the server.'
)
replace_paragraph_containing(
    'The security-specific test cases (MIT-01 and MIT-02) confirmed that:',
    'The security-specific verification confirmed that:'
)
replace_paragraph_containing(
    'All identified deficiencies in the test cases were resolved during the implementation iteration.',
    'During implementation, the authentication prototype was upgraded from a demo-only frontend flow to a database-backed flow with backend validation and password reset support.'
)
replace_paragraph_containing(
    'The testing phase confirms that security controls can be implemented at the user-interface level',
    'The testing phase confirms that security controls can be integrated into both the user interface and the backend service layer'
)

# Challenges / conclusion / future work
replace_paragraph_containing(
    'Demonstrating server-side validation (MIT-06) and HMAC payment verification (MIT-07) in a prototype without a real backend required careful design of the demonstration flow to make the security behavior visible.',
    'Integrating persistent credential storage while preserving the existing interface required careful alignment between frontend mutations, backend procedures, database schema, and validation rules.'
)
replace_paragraph_containing(
    'The FloraLink project successfully demonstrated the application of secure software development principles across the full project lifecycle.',
    'The FloraLink project successfully demonstrated the application of secure software development principles across the full project lifecycle. Starting from a structured requirements analysis that identified 22 functional use cases, 12 misuse cases, and 13 mitigation controls, the team produced a complete system design including UI prototypes, a 9-entity database schema, and a working authentication prototype that now includes database-backed registration, login, password reset, strong password validation, and temporary lockout behavior.'
)
replace_paragraph_containing(
    'Working prototype demonstrating UC-01, UC-02, MIT-01, and MIT-02 with a fully verified user interface.',
    'Working prototype demonstrating UC-01, UC-02, MIT-01, and MIT-02 with a verified browser flow and backend-backed account persistence, together with an added forgot-password/reset-password flow.'
)
replace_paragraph_containing(
    'Connect the prototype to a real persistent database (MySQL or PostgreSQL) with bcrypt-hashed credential storage.',
    'Extend the current database-backed prototype with additional business tables for products, carts, orders, and payments, and replace the direct reset form with token-based email recovery for production readiness.'
)

# Technology table (table index 63)
tech = doc.tables[63]
tech.rows[1].cells[0].text = 'Main programming language'
tech.rows[1].cells[1].text = 'TypeScript'
tech.rows[1].cells[2].text = 'Used across both the client and server for form logic, validation rules, database procedures, and application state.'

tech.rows[2].cells[0].text = 'User interface development'
tech.rows[2].cells[1].text = 'React + TSX/JSX'
tech.rows[2].cells[2].text = 'Used to build the registration, login, and reset-password screens and connect them to reactive state.'

tech.rows[3].cells[0].text = 'Client-server integration'
tech.rows[3].cells[1].text = 'tRPC + TanStack Query + SuperJSON'
tech.rows[3].cells[2].text = 'Used for typed frontend-to-backend mutations and query handling between the React interface and the Express server.'

tech.rows[4].cells[0].text = 'Interface styling'
tech.rows[4].cells[1].text = 'CSS3 + Tailwind CSS v4'
tech.rows[4].cells[2].text = 'Used for layout, spacing, typography, colors, states, and overall visual consistency.'

tech.rows[5].cells[0].text = 'Development and build tool'
tech.rows[5].cells[1].text = 'Vite + esbuild'
tech.rows[5].cells[2].text = 'Used for local development, frontend bundling, and production build generation.'

tech.rows[6].cells[0].text = 'Server runtime'
tech.rows[6].cells[1].text = 'Express + Node.js'
tech.rows[6].cells[2].text = 'Used to host the application, mount backend middleware, and expose the /api/trpc procedures.'

tech.rows[7].cells[0].text = 'Database and ORM'
tech.rows[7].cells[1].text = 'MySQL/TiDB + Drizzle ORM'
tech.rows[7].cells[2].text = 'Used to persist customer accounts, run schema migrations, and query/update authentication data.'

tech.rows[8].cells[0].text = 'Design, testing, and workflow tools'
tech.rows[8].cells[1].text = 'Figma, Canva, GitHub, VS Code, Vitest'
tech.rows[8].cells[2].text = 'Used for UI mockups, version control, development workflow, and automated backend route testing.'

# Login test table (table index 65)
login_tbl = doc.tables[65]
login_tbl.rows[1].cells[2].text = 'Correct registered email + correct password'
login_tbl.rows[1].cells[3].text = 'Login succeeds; welcome screen appears'
login_tbl.rows[3].cells[3].text = 'Error: invalid email or password displayed'

# Summary table (table index 68)
summary_tbl = doc.tables[68]
summary_tbl.rows[1].cells[2].text = 'Core flow passed'
summary_tbl.rows[1].cells[3].text = '0'
summary_tbl.rows[1].cells[4].text = 'Browser verification completed after DB integration'
summary_tbl.rows[2].cells[2].text = 'Core flow passed'
summary_tbl.rows[2].cells[3].text = '0'
summary_tbl.rows[2].cells[4].text = 'Browser verification completed after DB integration'
summary_tbl.rows[3].cells[2].text = 'Validated'
summary_tbl.rows[3].cells[3].text = '0'
summary_tbl.rows[3].cells[4].text = 'Rule enforcement confirmed in UI and backend schema validation'
summary_tbl.rows[4].cells[2].text = 'Validated'
summary_tbl.rows[4].cells[3].text = '0'
summary_tbl.rows[4].cells[4].text = 'Warning and lock-state behavior confirmed in the interface'
summary_tbl.rows[5].cells[2].text = 'Verified'
summary_tbl.rows[5].cells[3].text = '0'
summary_tbl.rows[5].cells[4].text = 'Supported by browser verification and automated backend route tests'

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(str(OUT))
