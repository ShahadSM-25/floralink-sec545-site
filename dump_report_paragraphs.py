from pathlib import Path
from docx import Document

DOC_PATH = Path('/home/ubuntu/upload/FloraLink_Final_Report_اساسي.docx')
doc = Document(str(DOC_PATH))
for i in range(430, 461):
    if i < len(doc.paragraphs):
        text = ' '.join(doc.paragraphs[i].text.split())
        print(f'P#{i}: {text}')
