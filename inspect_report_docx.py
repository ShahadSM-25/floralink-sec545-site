from pathlib import Path

DOC_PATH = Path('/home/ubuntu/upload/FloraLink_Final_Report_اساسي.docx')

try:
    from docx import Document
except ImportError:
    print('IMPORT_ERROR')
    raise

KEYWORDS = [
    'demo account',
    'in-memory',
    'Tailwind CSS',
    'Express static server',
    'CAPTCHA',
    'rate limiting',
    'lockout',
    'React + TSX/JSX',
    'database',
    'manual test',
    'Browser Developer Tools',
    'Chrome/Firefox',
    'stored demo account',
    'four implemented',
]

doc = Document(str(DOC_PATH))

print('PARAGRAPHS')
for i, p in enumerate(doc.paragraphs):
    text = ' '.join(p.text.split())
    if not text:
        continue
    lower = text.lower()
    if any(k.lower() in lower for k in KEYWORDS):
        print(f'P#{i}: {text}')

print('\nTABLE CELLS')
for ti, table in enumerate(doc.tables):
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            text = ' '.join(cell.text.split())
            if not text:
                continue
            lower = text.lower()
            if any(k.lower() in lower for k in KEYWORDS):
                print(f'T{ti} R{ri} C{ci}: {text}')
