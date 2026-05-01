from pathlib import Path
from docx import Document

DOC = Path('/home/ubuntu/floralink-sec545-site/FloraLink_Final_Report_corrected.docx')
doc = Document(str(DOC))

for p in doc.paragraphs:
    if p.text.strip() == 'The security-specific verification confirmed that:':
        p.text = ('The security-specific verification confirmed that weak passwords are rejected at the point of entry, '\
                  'common passwords are identified and blocked, the temporary lockout mechanism activates on the fifth failed attempt, '\
                  'and the countdown timer reflects the remaining lock duration correctly. Browser verification also confirmed that the '\
                  'database-backed password reset flow updates the stored credentials successfully and allows login with the new password.')
    if p.text.strip() == 'The testing phase confirms that security controls can be integrated into both the user interface and the backend service layer':
        p.text = ('The testing phase confirms that security controls can be integrated into both the user interface and the backend service layer in a way that improves security while preserving a usable experience. In the current FloraLink build, password policy enforcement is visible in the interface and revalidated in backend procedures, while account credentials are now persisted in the database rather than kept as demo-only frontend state. The black-box testing approach therefore remains appropriate because it evaluates the system from the perspective of a user and exposes the observable effects of both functional and security decisions.')

summary_tbl = doc.tables[68]
summary_tbl.rows[1].cells[2].text = '-'
summary_tbl.rows[1].cells[3].text = '-'
summary_tbl.rows[1].cells[4].text = 'Core browser flow verified after DB integration'
summary_tbl.rows[2].cells[2].text = '-'
summary_tbl.rows[2].cells[3].text = '-'
summary_tbl.rows[2].cells[4].text = 'Core browser flow verified after DB integration'
summary_tbl.rows[3].cells[2].text = '-'
summary_tbl.rows[3].cells[3].text = '-'
summary_tbl.rows[3].cells[4].text = 'Rules verified in UI and backend validation'
summary_tbl.rows[4].cells[2].text = '-'
summary_tbl.rows[4].cells[3].text = '-'
summary_tbl.rows[4].cells[4].text = 'Warning and lock state verified in interface'
summary_tbl.rows[5].cells[2].text = '-'
summary_tbl.rows[5].cells[3].text = '-'
summary_tbl.rows[5].cells[4].text = '24-case matrix retained; backend route tests passed'

doc.save(str(DOC))
print(str(DOC))
