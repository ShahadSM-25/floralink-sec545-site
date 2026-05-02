from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = Path('/home/ubuntu/floralink-sec545-site/FloraLink_Final_Report_submission_ready.docx')
OUT = Path('/home/ubuntu/floralink-sec545-site/FloraLink_Final_Report_submission_ready_v2.docx')


def insert_paragraph_after(paragraph, text='', style=None):
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def insert_table_after(paragraph, rows, cols, style=None):
    table = doc.add_table(rows=rows, cols=cols)
    if style:
        table.style = style
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def add_bullet_after(paragraph, text):
    p = insert_paragraph_after(paragraph, text)
    p.style = 'List Bullet'
    return p


doc = Document(str(SRC))

heading = None
first_future_bullet = None
for idx, p in enumerate(doc.paragraphs):
    txt = p.text.strip()
    if txt == '11.3 Future Work':
        heading = p
        # find first following non-empty paragraph
        for q in doc.paragraphs[idx + 1:]:
            if q.text.strip():
                first_future_bullet = q
                break
        break

if heading is None or first_future_bullet is None:
    raise RuntimeError('Could not locate Section 11.3 Future Work in the current report.')

# Rename existing heading and insert new content before future work bullets.
heading.text = '11.3 Security Testing Validation'

content_blocks = [
    ('paragraph', 'To strengthen the evaluation presented in the previous testing chapters, an additional security-focused validation activity was conducted for the authentication-related parts of FloraLink. This validation was designed to complement the documented black-box and backend route tests by checking how the implemented flows behaved when exposed to suspicious, malformed, and attack-like inputs. The focus remained on the currently implemented authentication scope, especially the registration, login, and password-reset use cases.'),
    ('paragraph', 'The security validation concentrated on three categories of checks: SQL injection-style input testing, fuzz testing, and low-intensity penetration-oriented negative testing. These activities were not intended to replace a full professional penetration-testing engagement. Instead, they were used to determine whether the implemented mitigations for input validation, strong password enforcement, reCAPTCHA protection, and temporary lockout behavior produced safe and predictable outcomes under adversarial conditions.'),
    ('heading2', '11.3.1 Scope of Security Validation'),
    ('paragraph', 'The validation targeted the externally reachable authentication request paths used by UC-01 and UC-02. Testing emphasized defensive behavior rather than exploitation. In other words, the goal was to confirm that hostile or abnormal inputs were rejected, contained, or handled as ordinary data without causing server crashes, information leakage, or unintended changes in authentication state.'),
    ('table_scope', None),
    ('heading2', '11.3.2 SQL Injection Testing'),
    ('paragraph', 'SQL injection-style payloads were supplied through user-controlled authentication fields in order to evaluate whether the backend would interpret attacker-controlled strings as executable query fragments. Classic SQL-like patterns were used in both the email and password fields. The observed behavior showed that structurally invalid payloads were rejected during validation, while suspicious strings placed in broader input fields were handled as ordinary text values. In the tested cases, no database error messages, stack traces, or abnormal authentication responses were exposed.'),
    ('paragraph', 'These results support the conclusion that the tested authentication routes resist basic SQL injection-style input attempts at the currently implemented layer. The evidence does not claim that every possible injection scenario has been eliminated from the wider FloraLink system, but it does show that the implemented authentication endpoints behave safely under the tested attack-like inputs.'),
    ('heading2', '11.3.3 Fuzz Testing'),
    ('paragraph', 'Fuzz testing was used to submit malformed, incomplete, oversized, and boundary-breaking inputs to the authentication flows. This included invalid email formats, too-short names, structurally incorrect request bodies, oversized text values, and weak passwords. The purpose of this activity was to determine whether FloraLink failed safely when receiving unexpected or abnormal input data.'),
    ('paragraph', 'The system responded with controlled validation errors in the tested scenarios. Invalid email patterns, missing or too-short fields, oversized values, and weak-password reset attempts were rejected without evidence of uncontrolled exceptions or unstable server behavior. This indicates that the current implementation applies meaningful server-side validation and input-boundary enforcement in the tested flows.'),
    ('heading2', '11.3.4 Penetration-Oriented Negative Testing'),
    ('paragraph', 'A smaller set of penetration-oriented negative tests was also performed to examine whether unauthenticated requests disclosed sensitive data, whether malformed request structures caused unstable behavior, and whether the authentication endpoints remained predictable under hostile but non-destructive conditions. These tests are better understood as focused security checks rather than as a full penetration-testing engagement.'),
    ('paragraph', 'The tested endpoints returned stable and controlled responses. Unauthenticated status checks did not reveal protected user information, and malformed request bodies produced explicit client-facing validation errors instead of internal server failures. Taken together, these outcomes provide additional confidence that the implemented authentication prototype remains robust when presented with unexpected request shapes and low-intensity offensive inputs.'),
    ('heading2', '11.3.5 Summary of Security Testing Cases'),
    ('table_cases', None),
    ('heading2', '11.3.6 Security Testing Conclusion'),
    ('paragraph', 'Overall, the additional security validation supports the broader testing conclusions already documented in this report. For the implemented authentication scope, FloraLink demonstrated sound behavior against malformed input, basic SQL injection-style payloads, and adversarial request-shape variations. These findings reinforce that the core mitigations visible in UC-01 and UC-02—namely input validation, reCAPTCHA integration, strong password enforcement, and temporary abuse-resistance controls—are reflected not only in design and implementation, but also in practical security-oriented testing outcomes.'),
    ('paragraph', 'Accordingly, this report classifies the completed activity as structured security testing of the implemented authentication prototype. It should be presented as a practical validation step that strengthens confidence in the current build, while a larger full-system penetration test remains appropriate future work for later project stages.'),
    ('heading2', '11.4 Future Work'),
]

# insert in reverse before first future bullet using insert_paragraph_before semantics
anchor = first_future_bullet

# helper to insert before anchor while preserving order by reverse iteration
for kind, payload in reversed(content_blocks):
    if kind == 'paragraph':
        p = anchor.insert_paragraph_before(payload)
        p.style = 'Normal'
        anchor = p
    elif kind == 'heading2':
        p = anchor.insert_paragraph_before(payload)
        p.style = 'Heading 2'
        anchor = p
    elif kind == 'table_scope':
        p = anchor.insert_paragraph_before('')
        table = insert_table_after(p, 1, 2, 'Table Grid')
        table.cell(0, 0).text = 'Validation Type'
        table.cell(0, 1).text = 'Purpose within FloraLink'
        rows = [
            ('SQL injection testing', 'To verify that attacker-controlled strings in authentication inputs are rejected safely or treated as ordinary data without triggering database errors.'),
            ('Fuzz testing', 'To verify that malformed, incomplete, and oversized inputs are handled through controlled validation rather than unstable execution.'),
            ('Penetration-oriented negative testing', 'To verify that authentication-related endpoints do not disclose sensitive information and remain stable under abnormal request structures.')
        ]
        for left, right in rows:
            row = table.add_row().cells
            row[0].text = left
            row[1].text = right
        anchor = p
    elif kind == 'table_cases':
        p = anchor.insert_paragraph_before('')
        table = insert_table_after(p, 1, 6, 'Table Grid')
        headers = ['Test Case ID', 'Module / Use Case', 'Input', 'Expected Output', 'Actual Output', 'Result']
        for i, h in enumerate(headers):
            table.cell(0, i).text = h
        rows = [
            ('SQL-01', 'UC-02 / Login', "Email field contains a classic SQL injection pattern such as ' OR 1=1 --", 'The request is rejected safely or treated as invalid input without query manipulation.', 'The request was rejected as invalid input and no database error information was exposed.', 'Pass'),
            ('SQL-02', 'UC-02 / Login', 'Password field contains SQL-like content while the remaining fields are syntactically valid.', 'The payload is processed as plain data and the system returns a normal authentication failure.', 'The system returned a standard invalid-credentials result without SQL error leakage.', 'Pass'),
            ('FZ-01', 'UC-01 / Registration', 'Malformed email, too-short full name, and structurally invalid registration fields.', 'Validation should reject malformed input before business logic is completed.', 'The endpoint returned controlled validation errors for the invalid fields.', 'Pass'),
            ('FZ-02', 'UC-01 / Registration', 'Oversized name and boundary-breaking field values.', 'Length validation should reject oversized values safely.', 'The endpoint rejected oversized values without crash or instability.', 'Pass'),
            ('FZ-03', 'UC-02 / Password Reset', 'Weak password input such as password.', 'Password policy should reject weak credentials.', 'The request was rejected according to password-strength validation rules.', 'Pass'),
            ('PT-01', 'Authentication State', 'Unauthenticated request to the current-user state endpoint.', 'No sensitive user data should be disclosed.', 'The endpoint returned no authenticated user data.', 'Pass'),
            ('PT-02', 'UC-02 / Login', 'Malformed JSON or invalid request-body structure.', 'The endpoint should fail safely with a controlled error.', 'The endpoint returned a controlled validation error rather than crashing.', 'Pass'),
        ]
        for row_values in rows:
            row = table.add_row().cells
            for i, value in enumerate(row_values):
                row[i].text = value
        anchor = p

# Update former future-work heading if still present elsewhere
for p in doc.paragraphs:
    if p is not heading and p.text.strip() == '11.3 Future Work':
        p.text = '11.4 Future Work'
        break

# If heading text isn't present separately because we inserted new 11.4, remove accidental duplicate by blanking a matching paragraph.
seen_future = 0
for p in doc.paragraphs:
    if p.text.strip() == '11.4 Future Work':
        seen_future += 1
        if seen_future == 1:
            continue
        # blank duplicates if any
        p.text = ''

# Clean any blank paragraph directly before the inserted 11.4 heading if present.
for i, p in enumerate(doc.paragraphs[:-1]):
    if p.text.strip() == '' and doc.paragraphs[i + 1].text.strip() == '11.4 Future Work':
        p.text = ''

# Save
OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(str(OUT))
